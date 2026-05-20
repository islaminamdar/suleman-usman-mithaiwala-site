#!/usr/bin/env python3
"""Assemble the manuscript into a single editable 6x9 DOCX (manuscript/book.docx).

The design matches Book One of the series, "The Dubai Syndicate Way":
Carlito type, a royal-purple / magenta palette, full-page chapter openers,
and proper title, copyright and contents pages.

Convert to PDF with LibreOffice:
  soffice --headless --convert-to pdf --outdir manuscript manuscript/book.docx
"""

import pathlib
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = pathlib.Path(__file__).resolve().parent.parent
MAN = ROOT / "manuscript"

CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

INK = RGBColor(0x1D, 0x1B, 0x19)
PURPLE = RGBColor(0x6B, 0x2C, 0x7E)
MAGENTA = RGBColor(0xB8, 0x3B, 0x7E)
MUTED = RGBColor(0x6F, 0x6A, 0x64)
RULE = "C7B2CF"
FONT = "Carlito"

BOOK_TITLE = "How to Start a Business in Dubai"
BOOK_SERIES = "The Dubai Syndicate Way"
BOOK_SUB = "The Practical Field Guide to 35 Industries"
BOOK_TAG = "From Restaurant to Real Estate, Salon to SaaS"
BOOK_AUTHOR = "Islam Inamdar"
BOOK_EDITION = "First Edition · 2026"
BOOK_PUBLISHER = "Published by Dubai Syndicate"

PARTS = [
    ("part-1-foundations", "Part One", "Foundations",
     "What every Dubai business shares — the city, the jurisdiction, "
     "the licence, the money, and the mindset."),
    ("part-2-industries", "Part Two", "The Field Guide",
     "Thirty-five industries, eight categories, one template — "
     "built to be read against each other."),
    ("part-3-resources", "Part Three", "Resources",
     "Directories, checklists, and tools to turn a decision "
     "into an application."),
]

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*)")
SUBTITLE = re.compile(r"^\*[^*].*[^*]\*$")
LISTITEM = re.compile(r"^([-*+]\s|\d+\.\s)")
H1_RE = re.compile(r"^#\s+(.+?)\s*$")
DOT = "·"

doc = Document()


def ordered_files():
    files = [p for p in MAN.rglob("*.md")
             if p.name not in ("OUTLINE.md", "BOOK.md")]
    return sorted(files, key=lambda p: str(p).lower())


def part_key(path):
    for key, *_ in PARTS:
        if key in str(path):
            return key
    return ""


def part_meta(key):
    return next(p for p in PARTS if p[0] == key)


def classify(path, h1):
    if path.name == "category-intro.md":
        return "category"
    if h1.startswith("Chapter "):
        return "chapter"
    if h1.startswith("Part "):
        return "partintro"
    return "resource"


def split_dot(text):
    if DOT in text:
        left, right = text.split(DOT, 1)
        return left.strip(), right.strip()
    return "", text.strip()


def parse_file(path):
    """Return (h1, subtitle, body_lines) — kicker line and H1 stripped out."""
    lines = path.read_text(encoding="utf-8").splitlines()
    h1, idx = "", 0
    for i, line in enumerate(lines):
        m = H1_RE.match(line.strip())
        if m:
            h1, idx = m.group(1).strip(), i + 1
            break
    rest = lines[idx:]
    while rest and not rest[0].strip():
        rest.pop(0)
    subtitle = ""
    if rest and SUBTITLE.match(rest[0].strip()) \
            and not rest[0].strip().startswith("**"):
        subtitle = rest[0].strip()[1:-1].strip()
        rest.pop(0)
    return h1, subtitle, rest


def is_special(s):
    return (not s or s.startswith("#") or s.startswith("|")
            or s.startswith(">") or re.fullmatch(r"[-*_]{3,}", s)
            or LISTITEM.match(s)
            or (SUBTITLE.match(s) and not s.startswith("**")))


# ---- low-level styling ---------------------------------------------------

def style_run(run, size=10.5, bold=False, italic=False, color=INK,
              tracking=0):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    if tracking:
        rpr = run._element.get_or_add_rPr()
        spc = OxmlElement("w:spacing")
        spc.set(qn("w:val"), str(tracking))
        rpr.insert_element_before(spc, "w:sz", "w:szCs", "w:u", "w:vertAlign",
                                  "w:rtl", "w:lang")


def add_runs(paragraph, text, size=10.5, color=INK):
    for tok in INLINE.split(text):
        if not tok:
            continue
        bold = italic = False
        if tok.startswith("**") and tok.endswith("**"):
            tok, bold = tok[2:-2], True
        elif tok.startswith("*") and tok.endswith("*"):
            tok, italic = tok[1:-1], True
        style_run(paragraph.add_run(tok), size=size, bold=bold,
                  italic=italic, color=color)


def para(text="", size=10.5, color=INK, align=None, italic_all=False,
         bold_all=False, space_before=0, space_after=6, indent=0.0,
         tracking=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.alignment = align if align is not None else JUSTIFY
    if indent:
        pf.left_indent = Inches(indent)
    if text:
        if italic_all or bold_all or tracking:
            style_run(p.add_run(text), size=size, color=color,
                      italic=italic_all, bold=bold_all, tracking=tracking)
        else:
            add_runs(p, text, size=size, color=color)
    return p


def add_borders(paragraph, top=False, bottom=False, size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side, enabled in (("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), RULE)
        pBdr.append(el)
    pPr.insert_element_before(pBdr, "w:tabs", "w:spacing", "w:ind",
                              "w:jc", "w:rPr", "w:sectPr")


def heading(text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = LEFT
    pf.keep_with_next = True
    if level == 2:
        pf.space_before = Pt(15)
        pf.space_after = Pt(5)
        style_run(p.add_run(text), size=12, bold=True, color=PURPLE)
    else:
        pf.space_before = Pt(11)
        pf.space_after = Pt(4)
        style_run(p.add_run(text), size=10.8, bold=True, color=INK)
    return p


def render_table(block):
    def cells(row):
        return [c.strip() for c in row.strip().strip("|").split("|")]
    header = cells(block[0])
    data = [cells(r) for r in block[2:]]
    ncol = len(header)
    table = doc.add_table(rows=1 + len(data), cols=ncol)
    table.style = "Table Grid"
    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = LEFT
        add_runs(cell.paragraphs[0], text, size=8.5, color=PURPLE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(data):
        for j in range(ncol):
            cell = table.rows[ri + 1].cells[j]
            cell.paragraphs[0].text = ""
            cell.paragraphs[0].alignment = LEFT
            add_runs(cell.paragraphs[0], row[j] if j < len(row) else "",
                     size=8.5)
    para(space_after=4)


def render_body(lines):
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("### "):
            heading(s[4:].strip(), 3)
            i += 1
            continue
        if s.startswith("## "):
            heading(s[3:].strip(), 2)
            i += 1
            continue
        if s.startswith("# "):
            heading(s[2:].strip(), 2)
            i += 1
            continue
        if re.fullmatch(r"[-*_]{3,}", s):
            i += 1
            continue
        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            if len(block) >= 2:
                render_table(block)
            continue
        m = re.match(r"^[-*+]\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, m.group(1))
            i += 1
            continue
        m = re.match(r"^\d+\.\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, m.group(1))
            i += 1
            continue
        if s.startswith(">"):
            para(re.sub(r"^>\s?", "", s), color=MUTED, align=LEFT,
                 italic_all=True, indent=0.3)
            i += 1
            continue
        buf = [s]
        i += 1
        while i < n and not is_special(lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        para(" ".join(buf))


# ---- structural pages ----------------------------------------------------

def render_title_page():
    for _ in range(4):
        para(space_after=0)
    para(BOOK_SERIES, size=11.5, color=MUTED, align=CENTER, space_after=30,
         tracking=46, bold_all=False)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = CENTER
    p.paragraph_format.space_after = Pt(18)
    style_run(p.add_run(BOOK_TITLE.upper()), size=27, bold=True,
              color=PURPLE)
    para(BOOK_SUB, size=14, color=INK, align=CENTER, space_after=5)
    para(BOOK_TAG, size=11, color=MUTED, align=CENTER, italic_all=True,
         space_after=44)
    para(BOOK_AUTHOR.upper(), size=12.5, color=INK, align=CENTER,
         tracking=34)


def render_copyright_page():
    doc.add_page_break()
    for _ in range(9):
        para(space_after=0)
    para(BOOK_TITLE, size=12, color=INK, align=CENTER, bold_all=True,
         space_after=8)
    para(f"Copyright © 2026 {BOOK_AUTHOR}", size=10, color=MUTED,
         align=CENTER, space_after=24)
    para("All rights reserved. No part of this publication may be "
         "reproduced, distributed, or transmitted in any form or by any "
         "means without the prior written permission of the author, "
         "except in the case of brief quotations embodied in critical "
         "reviews.", size=9.5, color=MUTED, align=CENTER, space_after=24)
    para(BOOK_EDITION, size=10, color=MUTED, align=CENTER, space_after=3)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = CENTER
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run("Book Two in "), size=10, color=MUTED)
    style_run(p.add_run(BOOK_SERIES), size=10, color=MUTED, italic=True)
    style_run(p.add_run(" series"), size=10, color=MUTED)
    para(BOOK_PUBLISHER, size=10, color=MUTED, align=CENTER)


def render_contents(files):
    doc.add_page_break()
    para(space_after=18)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = CENTER
    p.paragraph_format.space_after = Pt(22)
    style_run(p.add_run("CONTENTS"), size=19, bold=True, color=PURPLE,
              tracking=24)
    current = None
    for f in files:
        pk = part_key(f)
        if pk != current:
            current = pk
            _, label, name, _ = part_meta(pk)
            para(f"{label.upper()}  {DOT}  {name}", size=10, color=PURPLE,
                 align=LEFT, bold_all=True, space_before=14, space_after=6)
        h1, _, _ = parse_file(f)
        kind = classify(f, h1)
        if kind == "category":
            _, title = split_dot(h1)
            para(title, size=10, color=INK, align=LEFT, bold_all=True,
                 space_before=6, space_after=3, indent=0.16)
        elif kind == "chapter":
            left, title = split_dot(h1)
            m = re.search(r"\d+", left)
            num = m.group() if m else ""
            row = doc.add_paragraph()
            pf = row.paragraph_format
            pf.alignment = LEFT
            pf.left_indent = Inches(0.34)
            pf.space_after = Pt(2)
            pf.tab_stops.add_tab_stop(Inches(0.86))
            style_run(row.add_run(f"CH {num}\t"), size=10, bold=True,
                      color=PURPLE)
            style_run(row.add_run(title), size=10, color=INK)
        elif kind == "resource":
            para(h1, size=10, color=INK, align=LEFT, indent=0.34,
                 space_after=2)


def render_part_divider(key):
    _, label, name, blurb = part_meta(key)
    doc.add_page_break()
    for _ in range(8):
        para(space_after=0)
    para(label.upper(), size=12, color=MAGENTA, align=CENTER, space_after=14,
         bold_all=True, tracking=54)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = CENTER
    p.paragraph_format.space_after = Pt(14)
    style_run(p.add_run(name.upper()), size=24, bold=True, color=PURPLE)
    para(blurb, size=11, color=MUTED, align=CENTER, italic_all=True)


def render_opener(eyebrow, title, subtitle):
    doc.add_page_break()
    for _ in range(6):
        para(space_after=0)
    para(eyebrow.upper(), size=11, color=MAGENTA, align=CENTER,
         space_after=10, bold_all=True, tracking=44)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = CENTER
    pf.space_before = Pt(8)
    pf.space_after = Pt(10)
    pf.keep_with_next = True
    add_borders(p, top=True, bottom=True)
    style_run(p.add_run(title.upper()), size=21, bold=True, color=PURPLE)
    if subtitle:
        para(subtitle, size=11, color=MUTED, align=CENTER, italic_all=True,
             space_before=6)
    doc.add_page_break()


def render_category(kicker, title, subtitle, body_lines):
    doc.add_page_break()
    para(space_after=6)
    para(kicker.upper(), size=10.5, color=MAGENTA, align=CENTER,
         space_after=8, bold_all=True, tracking=44)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = CENTER
    p.paragraph_format.space_after = Pt(10)
    style_run(p.add_run(title.upper()), size=19, bold=True, color=PURPLE)
    if subtitle:
        para(subtitle, size=11, color=MUTED, align=CENTER, italic_all=True,
             space_after=18)
    render_body(body_lines)


def render_partintro(subtitle, body_lines):
    doc.add_page_break()
    para(space_after=6)
    if subtitle:
        para(subtitle, size=12, color=MUTED, align=CENTER, italic_all=True,
             space_after=20)
    render_body(body_lines)


def render_file(path):
    h1, subtitle, body = parse_file(path)
    kind = classify(path, h1)
    if kind == "chapter":
        left, title = split_dot(h1)
        m = re.search(r"\d+", left)
        num = int(m.group()) if m else 0
        render_opener(f"Chapter {num:02d}", title, subtitle)
        render_body(body)
    elif kind == "category":
        kicker, title = split_dot(h1)
        render_category(kicker, title, subtitle, body)
    elif kind == "partintro":
        render_partintro(subtitle, body)
    else:
        m = re.search(r"resource-(\d+)", path.name)
        num = int(m.group(1)) if m else 0
        render_opener(f"Resource {num:02d}", h1, subtitle)
        render_body(body)


def main():
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = section.right_margin = Inches(0.71)
    section.top_margin = section.bottom_margin = Inches(0.79)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.line_spacing = 1.42
    pf.space_after = Pt(6)

    render_title_page()
    render_copyright_page()
    files = ordered_files()
    render_contents(files)

    current = None
    for f in files:
        key = part_key(f)
        if key != current:
            current = key
            render_part_divider(key)
        render_file(f)

    doc.save(str(MAN / "book.docx"))
    print(f"Built manuscript/book.docx from {len(files)} files.")


if __name__ == "__main__":
    main()
